#!/usr/bin/env python3
"""Render project markdown manuscripts to docx with python-docx.

This script is intentionally lightweight but practical enough for internal
submission packages. It supports headings, paragraphs, inline emphasis,
editable inline/display equations, lists, fenced code blocks, images, and
GitHub-style pipe tables. Tables are formatted as three-line tables.
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable
import re
import textwrap

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import markdown

from equation_markup import equation_payload, omml_element, promote_inline_math


EMU_PER_INCH = 914400
TWIPS_PER_INCH = 1440


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_continuous_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def configure_document(doc: Document, chinese: bool = False, line_numbers: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    if line_numbers:
        add_continuous_line_numbers(section)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体" if chinese else "Times New Roman")

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if chinese else "Times New Roman")


def clear_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_edge(cell, edge: str, size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    element = tc_borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        tc_borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:color"), "000000")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_margins(cell, top: int = 45, start: int = 55, bottom: int = 45, end: int = 55) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)
    no_wrap.set(qn("w:val"), "true")


def _looks_numeric(value: str) -> bool:
    text = value.strip().lower()
    if not text:
        return False
    if text in {"na", "n/a", "none", "-"}:
        return True
    text = re.sub(
        r"\b(?:genes?|pairs?|studies?|comparisons?|seeds?|folds?)\b|基因|基因对|研究对|研究|比较",
        "",
        text,
    )
    text = text.replace("±", "").replace("%", "")
    text = re.sub(r"[\s,.;:()\[\]{}+\-/]", "", text)
    return bool(re.search(r"\d", text)) and not bool(re.search(r"[a-z\u4e00-\u9fff]", text))


def _percentile(values: list[int], fraction: float = 0.8) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    index = min(len(ordered) - 1, max(0, round((len(ordered) - 1) * fraction)))
    return float(ordered[index])


def _column_widths(parsed: list[list[str]], total_width: float) -> tuple[list[float], list[bool]]:
    column_count = max(len(row) for row in parsed)
    columns = [
        [row[index] if index < len(row) else "" for row in parsed]
        for index in range(column_count)
    ]
    numeric_columns: list[bool] = []
    demands: list[float] = []
    floors: list[float] = []

    if column_count <= 4:
        base_floor = 0.72
    elif column_count <= 6:
        base_floor = 0.52
    elif column_count <= 8:
        base_floor = 0.42
    elif column_count <= 10:
        base_floor = 0.35
    else:
        base_floor = 0.30

    for values in columns:
        body = [value for value in values[1:] if value.strip()]
        numeric = bool(body) and sum(_looks_numeric(value) for value in body) / len(body) >= 0.75
        numeric_columns.append(numeric)
        header_length = len(values[0]) if values else 0
        body_lengths = [min(len(value), 120) for value in body]
        body_p80 = _percentile(body_lengths)
        tokens = [
            token
            for value in values
            for token in re.split(r"[\s_+/\\-]+", value)
            if token
        ]
        longest_token = min(max((len(token) for token in tokens), default=0), 30)

        if numeric:
            demand = 0.48 + min(0.48, header_length * 0.014) + min(0.18, body_p80 * 0.010)
            floor = base_floor
        else:
            demand = 0.68 + min(1.70, body_p80 * 0.020) + min(0.48, longest_token * 0.016)
            demand = max(demand, 0.68 + min(1.0, header_length * 0.022))
            floor = base_floor * 1.35
        demands.append(demand)
        floors.append(floor)

    floor_total = sum(floors)
    if floor_total >= total_width:
        widths = [total_width / column_count] * column_count
        return widths, numeric_columns

    remaining = total_width - floor_total
    extra_demands = [max(0.01, demand - floor) for demand, floor in zip(demands, floors)]
    extra_total = sum(extra_demands)
    widths = [
        floor + remaining * extra / extra_total
        for floor, extra in zip(floors, extra_demands)
    ]
    return widths, numeric_columns


def _table_font_size(column_count: int, landscape: bool = False) -> float:
    if landscape:
        return 8.0 if column_count <= 10 else 7.6
    if column_count <= 3:
        return 9.2
    if column_count == 4:
        return 8.8
    if column_count <= 6:
        return 8.1
    if column_count <= 8:
        return 7.6
    if column_count <= 10:
        return 7.2
    return 7.0


def _wrap_friendly_table_text(value: str, column_count: int) -> str:
    if column_count < 7:
        return value
    return re.sub(r"([_/+])", lambda match: match.group(1) + "\u200b", value)


def format_table_layout(doc: Document, table, parsed: list[list[str]]) -> None:
    section = doc.sections[-1]
    available_width = (
        section.page_width - section.left_margin - section.right_margin
    ) / EMU_PER_INCH
    landscape = section.page_width > section.page_height
    total_width = min(9.35 if landscape else 6.85, float(available_width) - 0.05)
    column_count = len(table.columns)
    widths, numeric_columns = _column_widths(parsed, total_width)
    font_size = _table_font_size(column_count, landscape=landscape)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(round(total_width * TWIPS_PER_INCH)))

    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(round(width * TWIPS_PER_INCH)))

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            horizontal_margin = 35 if column_count >= 9 else 55
            set_cell_margins(cell, start=horizontal_margin, end=horizontal_margin)
            if row_index > 0 and numeric_columns[column_index]:
                set_cell_no_wrap(cell)
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(round(width * TWIPS_PER_INCH)))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if row_index == 0 or numeric_columns[column_index]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def add_table_section(doc: Document, landscape: bool) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = True
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def apply_three_line_table(table) -> None:
    rows = table.rows
    if not rows:
        return
    for row in rows:
        for cell in row.cells:
            clear_cell_borders(cell)
    set_repeat_table_header(rows[0])
    for cell in rows[0].cells:
        set_cell_edge(cell, "top")
        set_cell_edge(cell, "bottom")
    for cell in rows[-1].cells:
        set_cell_edge(cell, "bottom")
    for cell in rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def iter_inline(tag: Tag | NavigableString) -> Iterable[Tag | NavigableString]:
    if isinstance(tag, NavigableString):
        yield tag
    else:
        for child in tag.children:
            yield child


def add_runs(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    if node.name == "span" and "inline-math" in node.get("class", []):
        paragraph._p.append(omml_element(node.get("data-latex", "")))
        return

    if node.name == "a":
        text = node.get_text(" ", strip=False)
        href = node.get("href", "")
        run = paragraph.add_run(f"{text} ({href})" if href else text)
        run.underline = True
        return

    if node.name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        return

    text_only = node.get_text()
    if node.name in {"strong", "b"}:
        run = paragraph.add_run(text_only)
        run.bold = True
        return
    if node.name in {"em", "i"}:
        run = paragraph.add_run(text_only)
        run.italic = True
        return

    for child in node.children:
        add_runs(paragraph, child)


def paragraph_from_tag(doc: Document, tag: Tag, style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    for child in iter_inline(tag):
        add_runs(p, child)
    return p


def add_block_spacer(doc: Document, after_pt: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = Pt(1)


def add_equation(doc: Document, latex: str, number: str) -> None:
    """Insert a centered editable OMML equation with a right-side number."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (0.48, 5.89, 0.48)
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = Inches(width)
        clear_cell_borders(cell)
        set_cell_margins(cell, top=20, start=0, bottom=20, end=0)
        tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
        if tc_width is None:
            tc_width = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tc_width)
        tc_width.set(qn("w:type"), "dxa")
        tc_width.set(qn("w:w"), str(round(width * TWIPS_PER_INCH)))

    formula_paragraph = table.cell(0, 1).paragraphs[0]
    formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_paragraph.paragraph_format.space_before = Pt(3)
    formula_paragraph.paragraph_format.space_after = Pt(3)
    formula_paragraph._p.append(omml_element(latex))

    number_paragraph = table.cell(0, 2).paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.space_before = Pt(3)
    number_paragraph.paragraph_format.space_after = Pt(3)
    number_run = number_paragraph.add_run(f"({number})")
    number_run.font.name = "Times New Roman"
    number_run.font.size = Pt(11)

    add_block_spacer(doc, after_pt=3)


def add_image(doc: Document, src: str, base_dir: Path) -> None:
    image_path = (base_dir / src).resolve()
    if image_path.exists() and image_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp"}:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.0))
        return
    p = doc.add_paragraph(style="Intense Quote")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(f"[Figure file] {src}")


def handle_table(doc: Document, tag: Tag) -> None:
    rows = tag.find_all("tr")
    if not rows:
        return
    parsed = []
    max_cols = 0
    for row in rows:
        cells = row.find_all(["th", "td"])
        values = [cell.get_text(" ", strip=True) for cell in cells]
        max_cols = max(max_cols, len(values))
        parsed.append(values)
    use_landscape = max_cols >= 9
    if use_landscape:
        add_table_section(doc, landscape=True)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    for r_idx, values in enumerate(parsed):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            value = values[c_idx] if c_idx < len(values) else ""
            cell.text = _wrap_friendly_table_text(value, max_cols)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
    apply_three_line_table(table)
    format_table_layout(doc, table, parsed)
    add_block_spacer(doc, after_pt=8)
    if use_landscape:
        add_table_section(doc, landscape=False)


def handle_list(doc: Document, tag: Tag, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in tag.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        for child in li.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                continue
            add_runs(p, child)
        for child in li.children:
            if isinstance(child, Tag) and child.name == "ul":
                handle_list(doc, child, ordered=False)
            elif isinstance(child, Tag) and child.name == "ol":
                handle_list(doc, child, ordered=True)


def render_markdown(markdown_text: str, output_path: Path, asset_base_dir: Path) -> None:
    markdown_text = re.sub(r"\A---\n.*?\n---\n+", "", markdown_text, flags=re.S)
    markdown_text = textwrap.dedent(markdown_text)
    markdown_text = "\n".join(
        line[8:] if line.startswith("        ") else line
        for line in markdown_text.splitlines()
    )
    html = markdown.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "sane_lists"],
        output_format="html5",
    )
    soup = BeautifulSoup(html, "lxml")
    promote_inline_math(soup)
    doc = Document()
    configure_document(
        doc,
        chinese="_cn" in output_path.name,
        line_numbers="main_manuscript" in output_path.name,
    )
    base_dir = asset_base_dir

    root_nodes = soup.body.contents if soup.body else soup.contents
    for node in root_nodes:
        if isinstance(node, NavigableString):
            if str(node).strip():
                doc.add_paragraph(str(node).strip())
            continue
        if not isinstance(node, Tag):
            continue

        if node.name == "h1":
            paragraph_from_tag(doc, node, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
        elif node.name == "h2":
            paragraph_from_tag(doc, node, style="Heading 1")
        elif node.name == "h3":
            paragraph_from_tag(doc, node, style="Heading 2")
        elif node.name == "h4":
            paragraph_from_tag(doc, node, style="Heading 3")
        elif node.name == "h5":
            paragraph_from_tag(doc, node, style="Heading 4")
        elif node.name == "p":
            imgs = node.find_all("img", recursive=False)
            if imgs and len(imgs) == len([c for c in node.children if isinstance(c, Tag)]):
                for img in imgs:
                    add_image(doc, img.get("src", ""), base_dir)
            else:
                paragraph_from_tag(doc, node)
        elif node.name == "pre":
            payload = equation_payload(node)
            if payload is not None:
                add_equation(doc, *payload)
            else:
                p = doc.add_paragraph()
                run = p.add_run(node.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        elif node.name == "ul":
            handle_list(doc, node, ordered=False)
        elif node.name == "ol":
            handle_list(doc, node, ordered=True)
        elif node.name == "table":
            handle_table(doc, node)
        elif node.name == "blockquote":
            paragraph_from_tag(doc, node, style="Intense Quote")
        elif node.name == "hr":
            doc.add_paragraph()
        else:
            paragraph_from_tag(doc, node)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_markdown", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    markdown_text = args.input_markdown.read_text(encoding="utf-8")
    render_markdown(markdown_text, args.output_docx, args.input_markdown.parent)
    print(f"Wrote {args.output_docx}")


if __name__ == "__main__":
    main()
